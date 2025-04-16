import os
from flask import Flask, render_template, request, send_file
import fitz  # PyMuPDF
import unicodeconverter
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def extract_text_from_pdf(pdf_path):
    text_content = []
    doc = fitz.open(pdf_path)
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text:
            unicode_text = unicodeconverter.convert_bijoy_to_unicode(text)
            full_text = unicode_text + "\n"
            text_content.append(full_text)
        else:
            text_content.append(f"No text found.\n")
    doc.close()
    return text_content

def convert_bijoy_pdf_to_unicode_docx(pdf_path, output_docx_path):
    # Extract text with proper page separation
    text_content = extract_text_from_pdf(pdf_path)
    
    # Create and format Word document
    document = Document()
    
    # Add content page by page with proper formatting
    for page_num, page_text in enumerate(text_content):
        # Add page header
        document.add_heading(f"Page {page_num + 1}", level=2)
        
        # Add page content
        document.add_paragraph(page_text)
        
        # Add page break between pages
        document.add_page_break()
    
    document.save(output_docx_path)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'pdf_file' not in request.files:
            return "No file part"
        file = request.files['pdf_file']
        if file.filename == '':
            return "No selected file"
        if file:
            # Save uploaded file
            pdf_filename = file.filename
            pdf_path = os.path.join(UPLOAD_FOLDER, pdf_filename)
            file.save(pdf_path)

            # Create output .docx filename
            base_filename = os.path.splitext(pdf_filename)[0]
            docx_filename = base_filename + '.docx'
            docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)

            # Convert and return the file
            convert_bijoy_pdf_to_unicode_docx(pdf_path, docx_path)
            return send_file(docx_path, as_attachment=True)

    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)